"""
Evidence text extraction for AI Evidence (Next Layer).
Extracts text from PDF, DOCX, XLSX; stores result in EvidenceExtraction.
State: extract_pending → extracting → extracted | extract_failed.
"""
from __future__ import annotations

import io
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.models import EvidenceFile
from app.models.ai_evidence import EvidenceExtraction
from app.services import storage


# Content types we support for extraction
PDF_CT = "application/pdf"
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_CT = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
PNG_CT = "image/png"
JPEG_CT = "image/jpeg"


def _extract_pdf(data: bytes) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    full = "\n\n".join(parts).strip()
    return {
        "extracted_text": full,
        "has_text": bool(full),
        "has_tables": False,
        "page_count": len(reader.pages),
    }


def _extract_docx(data: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    full = "\n\n".join(parts).strip()
    tables = list(doc.tables)
    return {
        "extracted_text": full,
        "has_text": bool(full) or bool(tables),
        "has_tables": len(tables) > 0,
        "table_count": len(tables),
    }


def _extract_xlsx(data: bytes) -> dict:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            line = "\t".join(str(c) if c is not None else "" for c in row).strip()
            if line:
                parts.append(line)
    full = "\n".join(parts).strip()
    wb.close()
    return {
        "extracted_text": full,
        "has_text": bool(full),
        "has_tables": True,
    }


def _extract_image(_data: bytes) -> dict:
    # No OCR in MVP; placeholder for future Tesseract/cloud OCR
    return {
        "extracted_text": "",
        "has_text": False,
        "has_tables": False,
        "note": "image_no_ocr",
    }


def extract_text_from_bytes(data: bytes, content_type: str) -> dict:
    """Extract text/structure from file bytes. Returns dict for extraction_result."""
    ct = (content_type or "").split(";")[0].strip().lower()
    if ct == PDF_CT:
        return _extract_pdf(data)
    if ct == DOCX_CT:
        return _extract_docx(data)
    if ct == XLSX_CT:
        return _extract_xlsx(data)
    if ct in (PNG_CT, JPEG_CT):
        return _extract_image(data)
    return {
        "extracted_text": "",
        "has_text": False,
        "has_tables": False,
        "note": "unsupported_type",
        "content_type": content_type,
    }


async def get_or_create_extraction(
    db: AsyncSession,
    evidence_file_id: str,
    tenant_id: str,
    assessment_id: str,
    force_reextract: bool = False,
) -> tuple[EvidenceExtraction, bool]:
    """
    Get existing extraction for this evidence file (same assessment), or create new.
    Returns (extraction, created_new).
    """
    q = select(EvidenceExtraction).where(
        EvidenceExtraction.evidence_file_id == evidence_file_id,
        EvidenceExtraction.assessment_id == assessment_id,
        EvidenceExtraction.tenant_id == tenant_id,
    ).order_by(EvidenceExtraction.updated_at.desc()).limit(1)
    r = await db.execute(q)
    existing = r.scalar_one_or_none()
    if existing and not force_reextract:
        return existing, False
    if existing and force_reextract:
        existing.status = "extract_pending"
        existing.extraction_result = None
        existing.error_message = None
        await db.flush()
        return existing, False
    ext = EvidenceExtraction(
        tenant_id=tenant_id,
        assessment_id=assessment_id,
        evidence_file_id=evidence_file_id,
        status="extract_pending",
    )
    db.add(ext)
    await db.flush()
    return ext, True


async def run_extraction(db: AsyncSession, extraction_id: str) -> EvidenceExtraction | None:
    """
    Load extraction + evidence file, download file from storage, extract text, update extraction.
    Caller must commit. Returns updated EvidenceExtraction or None if not found.
    """
    q = select(EvidenceExtraction).where(EvidenceExtraction.id == extraction_id)
    r = await db.execute(q)
    ext = r.scalar_one_or_none()
    if not ext:
        return None
    qf = select(EvidenceFile).where(
        EvidenceFile.id == ext.evidence_file_id,
        EvidenceFile.tenant_id == ext.tenant_id,
    )
    rf = await db.execute(qf)
    evidence = rf.scalar_one_or_none()
    if not evidence:
        ext.status = "extract_failed"
        ext.error_message = "Evidence file not found"
        return ext
    ext.status = "extracting"
    await db.flush()
    try:
        data = storage.get_object_bytes(evidence.storage_key)
    except Exception as e:
        ext.status = "extract_failed"
        ext.error_message = str(e)[:500]
        ext.extraction_result = None
        return ext
    try:
        result = extract_text_from_bytes(data, evidence.content_type)
        ext.status = "extracted"
        ext.extraction_result = result
        ext.error_message = None
    except Exception as e:
        ext.status = "extract_failed"
        ext.error_message = str(e)[:500]
        ext.extraction_result = None
    return ext
